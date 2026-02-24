import asyncio
import os
from pathlib import Path
from typing import Dict, Any
from datetime import datetime

import pypdf
import docx

from .base import BaseParser, Document

class TextParser(BaseParser):
    """Parses standard text files (.txt, .md)."""
    
    async def parse(self, file_path: str) -> Document:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
            
        loop = asyncio.get_event_loop()
        
        # Async file read wrapper
        def read_file():
            with open(path, "r", encoding="utf-8") as f:
                return f.read()
                
        content = await loop.run_in_executor(None, read_file)
        
        metadata = {
            "source": str(path),
            "filename": path.name,
            "extension": path.suffix,
            "timestamp": datetime.fromtimestamp(path.stat().st_mtime).isoformat()
        }
        
        return Document(content=content, metadata=metadata)

class PDFParser(BaseParser):
    """Parses PDF files using pypdf."""
    
    async def parse(self, file_path: str) -> Document:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
            
        loop = asyncio.get_event_loop()
        
        def read_pdf():
            text_pages = []
            with open(path, "rb") as f:
                reader = pypdf.PdfReader(f)
                metadata = dict(reader.metadata) if reader.metadata else {}
                for i, page in enumerate(reader.pages):
                    text = page.extract_text()
                    if text:
                        text_pages.append(f"--- Page {i+1} ---\n{text}")
            return "\n\n".join(text_pages), metadata
            
        content, pdf_metadata = await loop.run_in_executor(None, read_pdf)
        
        metadata = {
            "source": str(path),
            "filename": path.name,
            "extension": path.suffix,
            "timestamp": datetime.fromtimestamp(path.stat().st_mtime).isoformat(),
            **pdf_metadata
        }
        
        return Document(content=content, metadata=metadata)

class DocxParser(BaseParser):
    """Parses DOCX files using python-docx."""
    
    async def parse(self, file_path: str) -> Document:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
            
        loop = asyncio.get_event_loop()
        
        def read_docx():
            doc = docx.Document(str(path))
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            return "\n".join(full_text)
            
        content = await loop.run_in_executor(None, read_docx)
        
        metadata = {
            "source": str(path),
            "filename": path.name,
            "extension": path.suffix,
            "timestamp": datetime.fromtimestamp(path.stat().st_mtime).isoformat()
        }
        
        return Document(content=content, metadata=metadata)

class ParserFactory:
    """Factory to get the correct parser based on file extension."""
    
    _parsers = {
        ".txt": TextParser(),
        ".md": TextParser(),
        ".pdf": PDFParser(),
        ".docx": DocxParser()
    }
    
    @classmethod
    def get_parser(cls, file_path: str) -> BaseParser:
        ext = Path(file_path).suffix.lower()
        if ext not in cls._parsers:
            raise ValueError(f"Unsupported file extension: {ext}")
        return cls._parsers[ext]
